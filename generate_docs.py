import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx not found. Please install it.")
    exit(1)

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Digital Twin Verse - Legal & Chargeback Prevention SOP', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Part 1: Terms & Conditions
    add_heading(doc, 'PART 1: TERMS & CONDITIONS (READY TO PUBLISH)', 1)
    
    doc.add_paragraph("Company Name: Digital Twin Verse\nProduct/Service: Digital Twin for Students (Ed-Tech Platform)\nJurisdiction: [City, State, India]\nEffective Date: [Insert Date]\n")
    
    add_heading(doc, '1. Acceptance of Terms', 2)
    doc.add_paragraph("By accessing or using the Digital Twin for Students platform (the \"Service\"), you agree to be bound by these Terms and Conditions. If you do not agree, do not use the Service. These terms apply to all students, parents, schools, and other users.")
    
    add_heading(doc, '2. Subscriptions, Payments & Pricing', 2)
    doc.add_paragraph("The Service is offered on a subscription and one-time purchase basis. Current price points are:\n- ₹119 for 6 months\n- ₹249 for 1 year\nAll payments are processed securely via our designated payment gateway, Razorpay. By providing payment information, you authorize Razorpay and Digital Twin Verse to charge the applicable fees.")
    
    add_heading(doc, '3. Refund and Cancellation Policy (STRICT)', 2)
    doc.add_paragraph("Digital Twin Verse offers a digital, non-tangible service. Due to the immediate access granted to premium features upon successful payment, ALL SALES ARE FINAL. We do not offer refunds, partial refunds, or credits for canceled subscriptions, except as expressly required by Indian law. If you wish to cancel future recurring billing, you must do so via your account dashboard before the next billing cycle.")
    
    add_heading(doc, '4. Mandatory First-Contact Dispute Resolution', 2)
    doc.add_paragraph("Before initiating a chargeback, refund request, or dispute through your bank or Razorpay, you MUST first contact our support team at [Support Email] to attempt an amicable resolution. We aim to resolve all legitimate service issues within 48 hours. Failing to contact us prior to initiating a bank dispute is a violation of these terms.")
    
    add_heading(doc, '5. Fraud, Misuse, and Account Termination', 2)
    doc.add_paragraph("Filing a false, fraudulent, or malicious payment dispute (e.g., claiming 'service not delivered' when login logs prove otherwise) constitutes a material breach of these terms and potential fraud. In the event of a false dispute:\na) Your account and access to the Service will be immediately and permanently terminated.\nb) We reserve the right to submit your IP address, device ID, login logs, and consent logs as evidence to Razorpay, your issuing bank, and relevant law enforcement cyber cells.\nc) We reserve the right to initiate legal proceedings under the Information Technology Act, 2000, and the Indian Penal Code to recover the disputed amount and legal costs.")
    
    add_heading(doc, '6. Consent-Logging Disclosure', 2)
    doc.add_paragraph("By making a payment, you explicitly acknowledge and agree that Digital Twin Verse actively logs your IP address, timestamp, device fingerprint, and checkbox consent state during the checkout process. This data is permanently stored and will be utilized strictly as evidence in the event of payment disputes or fraud investigations.")
    
    add_heading(doc, '7. Minors and Parental Consent', 2)
    doc.add_paragraph("The Service is intended for students. If you are under 18 years of age (a Minor), you must have the explicit consent of your parent or legal guardian to register and make payments. By completing a transaction, you represent that you are either a legal adult or possess verifiable parental consent. Digital Twin Verse does not intentionally collect data from minors without such consent, in compliance with the Digital Personal Data Protection Act, 2023 (DPDP).")
    
    add_heading(doc, '8. Grievance Officer', 2)
    doc.add_paragraph("In accordance with the Information Technology Act, 2000 and the Consumer Protection (E-Commerce) Rules, 2020, the name and contact details of the Grievance Officer are provided below:\nName: [Grievance Officer Name]\nEmail: [Grievance Email]\nPhone: [Phone Number]\nAddress: [Registered Office Address, City, State, India]")
    
    add_heading(doc, '9. Governing Law, Jurisdiction, and Arbitration', 2)
    doc.add_paragraph("These Terms shall be governed by the laws of India. Any dispute arising out of or in connection with these Terms, including false chargebacks, shall be subject to the exclusive jurisdiction of the courts in [City, State, India]. Alternatively, we reserve the right to resolve disputes via binding arbitration in [City, State] in English, under the Arbitration and Conciliation Act, 1996.")
    
    add_heading(doc, '10. Limitation of Liability and Indemnification', 2)
    doc.add_paragraph("To the maximum extent permitted by Indian law, Digital Twin Verse's total liability shall not exceed the amount you paid for the Service. You agree to indemnify and hold harmless Digital Twin Verse, its directors, and employees from any claims, damages, or costs (including legal fees) arising from your fraudulent use of the Service or false payment disputes.")
    
    doc.add_page_break()

    # Part 2: Fraud / Chargeback Prevention SOP
    add_heading(doc, 'PART 2: FRAUD / CHARGEBACK PREVENTION SOP', 1)
    
    add_heading(doc, 'Phase 1: Prevention (The Moat)', 2)
    doc.add_paragraph("1. Consent Capture: Implement the 'Consent-Gated Payment Popup' (provided in HTML). Do not allow the Razorpay UI to trigger until the user has checked all mandatory boxes.\n"
                      "2. Invoicing: Automatically email a GST-compliant invoice immediately upon successful payment. Ensure the email subject line is clear (e.g., 'Receipt for Digital Twin Premium - ₹249').\n"
                      "3. Welcome Email: Send a welcome email detailing exactly how to access the service. Include a clear 'Contact Support' button and explicitly state the 'No Refund' policy again.\n"
                      "4. Clear Descriptor: Ensure your Razorpay statement descriptor is highly recognizable (e.g., 'DIGITALTWINVERSE' not 'DTVLLP123').")

    add_heading(doc, 'Phase 2: Detection (Early Warning)', 2)
    doc.add_paragraph("1. Webhook Monitoring: Set up a backend listener for Razorpay's `dispute.created` webhook. This allows you to react instantly rather than waiting for a monthly report.\n"
                      "2. Chargeback Ratio Tracking: Keep your chargeback ratio below 1%. If you see a spike in disputes from a specific IP range or region, temporarily block that IP block.\n"
                      "3. Anomaly Detection: Flag accounts that upgrade to premium, download massive amounts of data, and immediately attempt to delete their account.")
    
    add_heading(doc, 'Phase 3: Response (The Evidence Package Template)', 2)
    doc.add_paragraph("When a dispute occurs, you have roughly 3-7 days to respond to Razorpay. Do not write custom emails. Use this standard Evidence Package:\n\n"
                      "DOCUMENTATION TO SUBMIT TO RAZORPAY:\n"
                      "a) Proof of Delivery: Backend logs showing the user's ID logging in after the payment timestamp, and the specific premium content they accessed.\n"
                      "b) Consent Log: A screenshot/JSON dump of the consent log (User ID, IP, Timestamp, 'Terms Accepted: True').\n"
                      "c) T&C Reference: A copy of Clause 3 (No Refunds) and Clause 4 (Mandatory First-Contact).\n"
                      "d) Invoice: The PDF invoice sent to the user.\n"
                      "e) Communication Log: Proof that the user never contacted support prior to the dispute (or if they did, the support ticket showing they were handled properly).")
    
    add_heading(doc, 'Phase 4: Escalation (Legal Action Against Malicious Users)', 2)
    doc.add_paragraph("If a user files a clearly fraudulent chargeback (e.g., they used the service for 5 months and then filed a 'Not Recognized' dispute with their bank) and the amount justifies the effort:\n"
                      "1. Send a Legal Notice: Send a formal legal notice via email and physical post (if address is known) citing fraud under Section 420 of the IPC and breach of contract.\n"
                      "2. Inform the Issuing Bank: Draft a formal letter to the grievance officer of the user's issuing bank, providing the evidence package and stating the user is committing friendly fraud.\n"
                      "3. Suspend Account & Ban Device: Immediately ban the user's email, IP address, and device fingerprint from future purchases.")

    doc.add_page_break()
    
    # Part 3: Lawyer Review List
    add_heading(doc, 'PART 3: WHAT NEEDS A LAWYER’S EYES', 1)
    doc.add_paragraph("To keep legal costs down, take this document to an Indian tech-lawyer and ask them to verify ONLY the following:\n\n"
                      "1. The Minors Data / Parental Consent Clause: Ensure it perfectly aligns with the latest enforcements of the DPDP Act 2023 for Ed-Tech platforms.\n"
                      "2. Refund Unenforceability: Verify if your specific digital good delivery exempts you entirely from mandatory cooling-off periods under the Consumer Protection (E-Commerce) Rules, 2020.\n"
                      "3. The Arbitration Venue: Confirm the enforceability of binding arbitration in your specific chosen city.")

    # Save
    doc_path = os.path.join(os.getcwd(), 'Digital_Twin_Verse_Legal_SOP.docx')
    doc.save(doc_path)
    print(f"Successfully generated {doc_path}")

if __name__ == "__main__":
    main()
