import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

add_heading('Digital Twin Verse - Legal & Compliance Package', 0)

# PART 1
add_heading('PART 1: FULL TERMS & CONDITIONS', 1)

doc.add_paragraph('''Last Updated: [DATE]
Company: Digital Twin Verse
Registered Office: [Samastipur, Bihar, India]
Email: [GRIEVANCE_OFFICER_EMAIL]
''')

add_heading('1. Acceptance of Terms', 2)
doc.add_paragraph('''By accessing or using the Digital Twin Verse platform ("Service"), you agree to be bound by these Terms & Conditions. If you do not agree, do not use the Service. If you are under the age of 18 ("Minor"), you must obtain verifiable consent from your parent or legal guardian before accessing premium features or making payments.''')

add_heading('2. Minors and Parental Consent', 2)
doc.add_paragraph('''Digital Twin Verse is an ed-tech platform. Users under 18 years of age may only make purchases or subscriptions with the explicit, verifiable consent of a parent or guardian. We reserve the right to verify parental consent via email or alternative methods. Parents may contact us at [CONTACT_EMAIL] to revoke consent or request deletion of a minor’s data.''')

add_heading('3. Refund and Cancellation Policy', 2)
doc.add_paragraph('''Digital Twin Verse provides immediate digital access to educational resources and platform features upon payment.

Non-Refundable Upon Access:
Except where strictly required by applicable law (including the Consumer Protection Act, 2019) or in cases of proven, documented technical failure on our end that prevents you from accessing the service, all subscription fees and one-time payments are non-refundable once the service has been accessed.

Cancellations:
You may cancel your subscription at any time to prevent future billing. However, cancellation does not entitle you to a refund for the current active billing cycle.

Technical Failures:
If you experience a severe technical failure preventing access, you must report it to [SUPPORT_EMAIL] within [48 HOURS] of the transaction. Refunds for technical failures are at the sole discretion of Digital Twin Verse after an internal investigation.''')

add_heading('4. Payment, Fraud Prevention, and Consent Logging', 2)
doc.add_paragraph('''By initiating a payment via Razorpay, you explicitly consent to our Fraud Prevention protocols. To protect against malicious chargebacks and disputes, we securely log your:
- IP Address
- Timestamp of transaction
- User Agent (Browser/Device data)
- Exact acceptance of these Terms and our No-Refund policy.

Account Termination for Fraud:
Filing a false or malicious chargeback or payment dispute after legitimately receiving access to the Service constitutes a breach of these Terms. Digital Twin Verse reserves the right to immediately terminate the account, permanently ban the user, and use the logged consent data as evidence in dispute arbitration with Razorpay or legal authorities.''')

add_heading('5. Mandatory First-Contact Dispute Resolution', 2)
doc.add_paragraph('''Before initiating any chargeback with your bank or payment gateway, you agree to contact our support team at [SUPPORT_EMAIL] to attempt to resolve the issue in good faith. Failure to contact us prior to filing a dispute may be submitted as evidence of bad faith during the chargeback arbitration process.''')

add_heading('6. Limitation of Liability & Indemnification', 2)
doc.add_paragraph('''To the maximum extent permitted by Indian law, Digital Twin Verse shall not be liable for any indirect, incidental, or consequential damages arising out of your use of the Service. You agree to indemnify and hold harmless Digital Twin Verse, its officers, and employees from any claims, damages, or expenses arising from your violation of these Terms or applicable laws.''')

add_heading('7. Governing Law and Arbitration', 2)
doc.add_paragraph('''These Terms shall be governed by and construed in accordance with the laws of India. Any dispute arising out of or relating to these Terms, which cannot be resolved amicably, shall be resolved through binding arbitration under the Arbitration and Conciliation Act, 1996. The seat of arbitration shall be [Samastipur, Bihar, India].''')

add_heading('8. Grievance Officer', 2)
doc.add_paragraph('''In compliance with the Information Technology Act, 2000 and the Consumer Protection (E-Commerce) Rules, 2020, the Grievance Officer for Digital Twin Verse is:

Name: [GRIEVANCE_OFFICER_NAME]
Email: [GRIEVANCE_OFFICER_EMAIL]
Address: [FULL_REGISTERED_ADDRESS, Samastipur, Bihar]

The Grievance Officer will acknowledge receipt of any consumer complaint within 48 hours and will redress the complaint within one month from the date of receipt.''')


# PART 2
doc.add_page_break()
add_heading('PART 2: MINOR / PARENTAL CONSENT MECHANISM', 1)
doc.add_paragraph('''Why a single self-tick checkbox is legally weak:
Under the Indian Contract Act, 1872, minors (under 18) are generally incompetent to contract. If a minor clicks "I agree," the contract is void ab initio (invalid from the start). A simple checkbox does not prove the person clicking is an adult. In a chargeback dispute, parents can claim "my child bought this without my permission," and you will almost certainly lose.

Improved Flow (High-Level Design for Ed-Tech):
1. Age Gate at Signup: Ask the user for their Date of Birth during account creation (not just at payment).
2. Parent Verification (If Minor): If the DOB indicates they are under 18, require a Parent/Guardian Email Address.
3. Transaction Hold: When the minor attempts to purchase a subscription, send an automated email to the Parent/Guardian Email with a secure, one-time link.
4. Parental Authorization: The parent clicks the link, views the payment details and terms, and clicks "I Authorize this Purchase." This step captures the Parent's IP and Timestamp.
5. Payment Unlock: Once authorized, the minor's account is unlocked to proceed to the Razorpay checkout.

Immediate Alternative (Friction-Balanced):
If full email verification causes too much drop-off, implement an "Adult Verification Checkbox" explicitly worded: "I confirm that I am over 18 years of age OR I am the parent/guardian authorizing this transaction for a minor." Log the IP and user session. Combine this with the backend Consent Log. While not bulletproof against a determined parent, it shifts the burden of proof in Razorpay disputes by demonstrating the user actively falsified their identity.''')


# PART 3
doc.add_page_break()
add_heading('PART 3: FRAUD / CHARGEBACK PREVENTION SOP', 1)
doc.add_paragraph('''1. PREVENTION Phase:
- Mandatory Consent Gate: Implement the Consent-Gated Popup (Part 4) before every Razorpay transaction. No bypass allowed.
- Invoice Generation: Automatically email a detailed receipt to the user post-purchase detailing the plan, access links, and a reminder of the no-refund policy.
- Account Activity Logging: Track the user's first login post-purchase. (If they log in and consume content, it proves "Service Delivered").

2. DETECTION Phase:
- Webhook Monitoring: Set up Razorpay webhooks for `chargeback.created` or `dispute.created`.
- Alerting: Route webhook alerts immediately to an internal Slack channel or Admin Email so disputes are caught on Day 1.

3. RESPONSE Phase (Evidence Package Template):
When a dispute arrives, submit a single PDF to Razorpay containing:
A. Proof of Service Delivery: Server logs showing the user logged in and accessed the course material AFTER the payment timestamp.
B. Proof of Consent: Database extract showing the user's ID, the exact Timestamp they clicked "Accept", their IP Address, and the specific terms they agreed to.
C. Proof of Communication: Screenshot of the automated receipt/welcome email sent to their registered email.

4. ESCALATION Phase (Legal Notice):
If a user files a malicious dispute (e.g., claiming "Fraud" when logs show they completed 3 courses), send a formal Cease & Desist / Legal Notice via email citing Section 415 (Cheating) of the Indian Penal Code. Offer them 48 hours to withdraw the dispute with the bank before escalating.''')


# PART 4
doc.add_page_break()
add_heading('PART 4: CONSENT-GATED PAYMENT POPUP (REFERENCE)', 1)
doc.add_paragraph('''This popup must appear when the user clicks "Buy Now / Upgrade" and MUST block the Razorpay checkout until completed.

Key Features:
- Short, readable summary (no legal jargon).
- Safe refund wording.
- Scroll-to-bottom enforcement.
- Mandatory checkboxes.
- Captures IP, Timestamp, and User ID via AJAX to your backend before unlocking payment.

*See the accompanying standalone HTML file for the actual code implementation.*''')

# FINAL REVIEW NOTES
doc.add_page_break()
add_heading('FINAL REVIEW ITEMS FOR YOUR LAWYER', 1)
doc.add_paragraph('''1. Refund Clause Enforceability: Ensure the specific "technical failure" carve-out aligns with the exact nature of your digital delivery mechanism under the CP Act 2019.
2. Minor-Consent Mechanism: Verify if your specific implementation of parent email verification satisfies the emerging standards of the Digital Personal Data Protection (DPDP) Act for processing children's data.
3. Arbitration Clause: Confirm that [Samastipur, Bihar] is the ideal strategic seat for arbitration, considering the location of your legal counsel and operations.''')

doc.save('Digital_Twin_Verse_Legal_SOP_V2.docx')
print("Document generated successfully.")
